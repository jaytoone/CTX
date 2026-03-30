"""
CTX BM25 교체 후 성과 보고서 생성기
산출물: docs/research/CTX_BM25_PERFORMANCE_REPORT.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 팔레트 ────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_BLUE   = RGBColor(0x16, 0x21, 0x3E)
C_ACCENT = RGBColor(0x0F, 0x3D, 0x5C)
C_GREEN  = RGBColor(0x1B, 0x5E, 0x20)
C_RED    = RGBColor(0x7B, 0x1F, 0x1F)
C_GOLD   = RGBColor(0xB8, 0x86, 0x0B)
C_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
C_LGRAY  = RGBColor(0xE8, 0xEC, 0xF0)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT   = RGBColor(0x1C, 0x1C, 0x1C)
C_PALE_G = RGBColor(0xE8, 0xF5, 0xE9)
C_PALE_R = RGBColor(0xFF, 0xEB, 0xEE)
C_PALE_Y = RGBColor(0xFF, 0xF9, 0xC4)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_font(cell, text, bold=False, size=10, color=C_TEXT,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_section_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = C_BLUE if level == 1 else C_ACCENT
    return p


def add_body(doc, text, italic=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        if italic:
            run.font.italic = True
    return p


def build_report():
    doc = Document()

    # ── 페이지 여백 ──────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # ── 표지 ────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_run = title_p.add_run("CTX 성과 보고서 — BM25 교체 후")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = C_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Document Retrieval Benchmark v2 | 29 docs · 87 queries · 2026-03-27")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── 1. 전략별 전체 지표 ──────────────────────────────────
    add_section_title(doc, "1. 전략별 전체 지표 (87 queries)")
    add_body(doc, "CTX-doc은 BM25와 Dense TF-IDF 기준선을 모든 전체 지표에서 초과합니다.")

    overall_headers = ["전략", "R@3", "R@5", "NDCG@5", "MRR", "평가"]
    overall_rows = [
        ["CTX-doc (heading+BM25)", "0.839", "0.931", "0.787", "0.748", "✅ 1위"],
        ["BM25 단독",              "0.667", "0.839", "0.655", "0.611", "—"],
        ["Dense TF-IDF",           "0.690", "0.805", "0.607", "0.563", "—"],
    ]
    col_widths = [2.2, 0.7, 0.7, 0.8, 0.7, 0.9]
    t1 = doc.add_table(rows=1 + len(overall_rows), cols=len(overall_headers))
    t1.style = 'Table Grid'

    # 헤더
    for j, h in enumerate(overall_headers):
        cell = t1.rows[0].cells[j]
        cell.width = Inches(col_widths[j])
        set_cell_bg(cell, C_BLUE)
        set_cell_font(cell, h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    row_bgs = [C_PALE_G, C_LGRAY, C_LGRAY]
    for i, row_data in enumerate(overall_rows):
        row = t1.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, row_bgs[i])
            bold = (i == 0)
            align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(cell, val, bold=bold, align=align)

    doc.add_paragraph()

    # ── 2. 쿼리 타입별 R@3 ──────────────────────────────────
    add_section_title(doc, "2. 쿼리 타입별 R@3 비교")
    add_body(doc, "CTX는 heading 기반 쿼리에서 압도적이나, keyword 쿼리에서 BM25 단독 대비 -6.9%p 열세.")

    type_headers = ["쿼리 타입", "CTX R@3", "BM25 단독 R@3", "Delta", "평가"]
    type_rows = [
        ["heading_paraphrase", "1.000", "0.655", "+34.5%p", "✅ 완벽"],
        ["heading_exact",      "0.862", "0.621", "+24.1%p", "✅ 우위"],
        ["keyword",            "0.655", "0.724", "-6.9%p",  "⚠️ 열세"],
    ]
    type_col_w = [2.0, 1.0, 1.2, 1.0, 1.0]
    row_bgs2 = [C_PALE_G, C_PALE_G, C_PALE_Y]

    t2 = doc.add_table(rows=1 + len(type_rows), cols=len(type_headers))
    t2.style = 'Table Grid'
    for j, h in enumerate(type_headers):
        cell = t2.rows[0].cells[j]
        set_cell_bg(cell, C_ACCENT)
        set_cell_font(cell, h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row_data in enumerate(type_rows):
        row = t2.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, row_bgs2[i])
            delta_color = C_RED if val.startswith('-') else C_TEXT
            color = delta_color if j == 3 else C_TEXT
            align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(cell, val, color=color, align=align)

    doc.add_paragraph()

    # ── 3. BM25 교체 전후 비교 ──────────────────────────────
    add_section_title(doc, "3. TF-IDF → BM25 교체 전후 비교")

    before_after_headers = ["지표", "교체 전 (TF-IDF)", "교체 후 (BM25)", "개선폭", "목표"]
    before_after_rows = [
        ["keyword R@3",          "0.379", "0.655", "+27.6%p", "≥ 0.600 ✅"],
        ["heading_paraphrase R@3","0.966", "1.000", "+3.4%p",  "—"],
        ["heading_exact R@3",    "0.793", "0.862", "+6.9%p",  "—"],
        ["전체 R@3",              "0.713", "0.839", "+12.6%p", "—"],
        ["전체 R@5",              "0.862", "0.931", "+6.9%p",  "—"],
        ["NDCG@5",                "0.717", "0.787", "+7.0%p",  "—"],
        ["MRR",                   "0.688", "0.748", "+6.0%p",  "—"],
    ]
    ba_col_w = [1.8, 1.2, 1.2, 1.0, 1.0]

    t3 = doc.add_table(rows=1 + len(before_after_rows), cols=len(before_after_headers))
    t3.style = 'Table Grid'
    for j, h in enumerate(before_after_headers):
        cell = t3.rows[0].cells[j]
        set_cell_bg(cell, C_DARK)
        set_cell_font(cell, h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row_data in enumerate(before_after_rows):
        row = t3.rows[i + 1]
        bg = C_PALE_G if i == 0 else (C_LGRAY if i % 2 == 0 else C_WHITE)
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            bold = (j == 2)
            set_cell_font(cell, val, bold=bold, align=align)

    doc.add_paragraph()

    # ── 4. 코드 검색 vs Nemotron ────────────────────────────
    add_section_title(doc, "4. 코드 검색 성과 (CTX vs Nemotron-Cascade-2)")
    add_body(doc, "이전 실험 (166 queries, synthetic code retrieval benchmark) 결과.")

    code_headers = ["지표", "CTX", "Nemotron", "Delta", "p-value", "평가"]
    code_rows = [
        ["R@5",  "0.958", "0.946", "+1.2%p",   "0.629 (ns)", "동등"],
        ["TES",  "0.668", "0.241", "+177%",     "3.1e-27 ***","✅ 압도"],
        ["Cross-session R@10","0.567","—",       "독점",      "—",    "✅ LLM 불가"],
    ]
    code_col_w = [1.8, 0.8, 0.9, 0.9, 1.1, 0.9]
    row_bgs3 = [C_LGRAY, C_PALE_G, C_PALE_G]

    t4 = doc.add_table(rows=1 + len(code_rows), cols=len(code_headers))
    t4.style = 'Table Grid'
    for j, h in enumerate(code_headers):
        cell = t4.rows[0].cells[j]
        set_cell_bg(cell, C_ACCENT)
        set_cell_font(cell, h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row_data in enumerate(code_rows):
        row = t4.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, row_bgs3[i])
            align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(cell, val, align=align)

    doc.add_paragraph()

    # ── 5. 잘하는 것 vs 못하는 것 ───────────────────────────
    add_section_title(doc, "5. 강점 / 약점 요약")

    sw_headers = ["분류", "항목", "수치", "근거"]
    sw_rows = [
        ["✅ 강점", "Heading paraphrase 문서 서페이싱", "R@3 = 1.000", "concept_index 역인덱스 완벽 매칭"],
        ["✅ 강점", "토큰 효율성 (코드 검색)",          "TES = 0.668", "Nemotron 0.241 대비 +177%"],
        ["✅ 강점", "크로스세션 파일 복원",              "R@10 = 0.567","persistent_memory + SessionStart hook"],
        ["✅ 강점", "코드 R@5 (LLM 방식 동등)",         "0.958 ≈ Nem", "p=0.629 (통계적 비차이)"],
        ["⚠️ 열세", "Keyword 쿼리 문서 검색",           "R@3 = 0.655", "BM25 단독 0.724 대비 -6.9%p"],
        ["❌ 약점", "외부 코드베이스 일반화",            "R@5 = 0.152", "내부 대비 85% 하락, heuristic 과적합"],
        ["❌ 약점", "교차 파일 추론",                   "Hop2 미지원", "import graph 2-hop 한계"],
    ]
    sw_col_w = [0.8, 2.0, 1.2, 2.2]
    row_bgs4 = [C_PALE_G]*4 + [C_PALE_Y] + [C_PALE_R]*2

    t5 = doc.add_table(rows=1 + len(sw_rows), cols=len(sw_headers))
    t5.style = 'Table Grid'
    for j, h in enumerate(sw_headers):
        cell = t5.rows[0].cells[j]
        set_cell_bg(cell, C_DARK)
        set_cell_font(cell, h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row_data in enumerate(sw_rows):
        row = t5.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, row_bgs4[i])
            bold = (j == 0)
            set_cell_font(cell, val, bold=bold)

    doc.add_paragraph()

    # ── 6. 지표 설명 ────────────────────────────────────────
    add_section_title(doc, "6. 지표 설명")

    metric_headers = ["지표", "의미", "범위", "해석 기준"]
    metric_rows = [
        ["R@k (Recall@k)",
         "상위 k개 결과 안에 정답이 있는 비율",
         "0 ~ 1",
         "k=3: 3번 안에 찾냐? 0.8+ = 실용 수준"],
        ["NDCG@5",
         "정답 순위에 로그 감소 가중치 부여 (rank 1이 5보다 훨씬 높은 점수)",
         "0 ~ 1",
         "0.7+ = 높은 품질"],
        ["MRR (Mean Reciprocal Rank)",
         "각 쿼리에서 정답이 첫 등장한 순위의 역수 평균",
         "0 ~ 1",
         "0.5 = 평균 2등, 0.7+ = 우수"],
        ["TES (Token Efficiency Score)",
         "정답 포함에 사용한 토큰 비율의 역수 — 적은 토큰으로 정답 포함할수록 높음",
         "0 ~ 1",
         "0.5+ = 전체 컨텍스트 대비 50% 절감"],
    ]
    me_col_w = [1.3, 2.5, 0.7, 1.7]

    t6 = doc.add_table(rows=1 + len(metric_rows), cols=len(metric_headers))
    t6.style = 'Table Grid'
    for j, h in enumerate(metric_headers):
        cell = t6.rows[0].cells[j]
        set_cell_bg(cell, C_ACCENT)
        set_cell_font(cell, h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row_data in enumerate(metric_rows):
        row = t6.rows[i + 1]
        bg = C_LGRAY if i % 2 == 0 else C_WHITE
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            bold = (j == 0)
            set_cell_font(cell, val, bold=bold, size=9)

    # ── 푸터 ────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "CTX BM25 Performance Report | omc-live iter 3 | 2026-03-27"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
    footer_run.font.italic = True

    out_path = Path(__file__).parent / "CTX_BM25_PERFORMANCE_REPORT.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


print(build_report())
