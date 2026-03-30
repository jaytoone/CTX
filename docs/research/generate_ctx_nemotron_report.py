"""
CTX vs Nemotron-Cascade-2 최종 비교 보고서 생성기
산출물: CTX_NEMOTRON_COMPARISON_REPORT.docx
"""
import json
import math
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 팔레트 ────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_BLUE   = RGBColor(0x16, 0x21, 0x3E)
C_ACCENT = RGBColor(0x0F, 0x3D, 0x5C)
C_GREEN  = RGBColor(0x1B, 0x5E, 0x20)
C_RED    = RGBColor(0x7B, 0x1F, 0x1F)
C_GOLD   = RGBColor(0xB8, 0x86, 0x0B)
C_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
C_LGRAY  = RGBColor(0xE8, 0xEC, 0xF0)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT   = RGBColor(0x1C, 0x1C, 0x1C)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_font(cell, text, bold=False, size=10, color=C_TEXT,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_section_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = C_BLUE if level == 1 else C_ACCENT
    return p


def add_body(doc, text, italic=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        if italic:
            run.font.italic = True
    return p


# ── 데이터 ────────────────────────────────────────────────────
OVERALL_METRICS = [
    {
        "metric": "Recall@1",
        "ctx": "0.688",
        "nem": "0.528",
        "delta": "+0.160",
        "p_value": "2.2e-4 ***",
        "effect_d": "0.336 (medium)",
        "winner": "CTX",
        "description": (
            "첫 번째 결과가 관련 파일일 확률. CTX는 트리거 기반 심볼 인덱스로 "
            "1순위 정확도가 높음. Nemotron은 5개 중 어딘가에 포함되지만 1순위 배치가 약함."
        ),
    },
    {
        "metric": "Recall@5",
        "ctx": "0.958",
        "nem": "0.946",
        "delta": "+0.012",
        "p_value": "0.629 ns",
        "effect_d": "0.042 (negligible)",
        "winner": "동등",
        "description": (
            "상위 5개 결과 안에 관련 파일이 포함될 확률. "
            "두 시스템 모두 약 95% 이상으로 통계적으로 동등(p=0.629). "
            "커버리지 측면에서 실질적 차이 없음."
        ),
    },
    {
        "metric": "NDCG@5",
        "ctx": "0.929",
        "nem": "0.850",
        "delta": "+0.079",
        "p_value": "1.4e-4 ***",
        "effect_d": "0.191 (small)",
        "winner": "CTX",
        "description": (
            "Normalized Discounted Cumulative Gain — 상위 순위에 관련 파일이 "
            "배치될수록 점수 증가. CTX의 구조적 우선순위 지정이 Nemotron의 "
            "LLM 유사도 기반 랭킹보다 정밀함."
        ),
    },
    {
        "metric": "TES",
        "ctx": "0.668",
        "nem": "0.241",
        "delta": "+0.428",
        "p_value": "3.1e-27 ***",
        "effect_d": "1.322 (large)",
        "winner": "CTX",
        "description": (
            "Trade-off Efficiency Score = Recall@5 / log(1 + files_loaded). "
            "CTX는 평균 ~4.4개 파일만 로드(실측), Nemotron은 전체 50개 로드. "
            "CTX가 2.78× 효율적. 가장 큰 격차 (대형 효과 크기 d=1.322)."
        ),
    },
    {
        "metric": "Token Efficiency",
        "ctx": "0.098",
        "nem": "1.000",
        "delta": "−0.902",
        "p_value": "—",
        "effect_d": "—",
        "winner": "CTX",
        "description": (
            "사용 토큰 / 전체 코드베이스 토큰. CTX는 9.8%만 사용(선택적 로드), "
            "Nemotron은 100% 사용(전체 컨텍스트 방식). "
            "실제 비용: Nemotron이 10× 더 많은 토큰 소비."
        ),
    },
]

TRIGGER_METRICS = [
    ("EXPLICIT_SYMBOL", 79, "0.911", "0.962", "+0.051", "0.206 ns",
     "함수명·클래스명 직접 검색. Nemotron이 코드 의미 이해로 +5pp 우위 (통계적 비유의)"),
    ("SEMANTIC_CONCEPT", 72, "1.000", "0.946", "−0.054", "0.011 *",
     "개념 수준 검색 ('캐싱 관련 코드'). CTX 트리거 분류기가 모듈 목적 이해로 완벽 달성"),
    ("TEMPORAL_HISTORY", 10, "1.000", "0.900", "−0.100", "1.000 ns",
     "시간적 변경 이력 기반 검색. CTX는 git blame/history 메타데이터 활용 가능"),
    ("IMPLICIT_CONTEXT", 5, "1.000", "0.767", "−0.233", "0.250 ns",
     "명시적 언급 없이 문맥으로 파일 추론. CTX 구조적 규칙셋이 LLM 어휘 유사도보다 우수"),
]

MULTI_SYSTEM = [
    ("Full Context (unranked)", "0.014", "0.075", "0.052", "0.019", "100%",
     "모든 파일 동점 반환 — 관련 파일이 상위 5위 안에 들어올 확률이 낮음"),
    ("GraphRAG", "0.318", "0.514", "0.415", "0.214", "22.5%",
     "그래프 구조 기반 검색 — 중간 수준, 임포트 의존성 반영"),
    ("RANGER", "0.318", "0.345", "0.342", "0.249", "5.8%",
     "희소 근사 — 토큰 효율은 좋으나 recall 낮음"),
    ("Nemotron (LLM ranking)", "0.528", "0.946", "0.850", "0.241", "100%",
     "전체 컨텍스트 LLM 랭킹 — 높은 recall, 10× 토큰 비용"),
    ("BM25", "0.745", "0.982", "0.960", "0.410", "18.7%",
     "어휘 매칭 — Recall@5 최고, Nemotron보다 낮은 비용으로 더 높은 recall"),
    ("CTX (adaptive_trigger)", "0.688", "0.958", "0.929", "0.668", "9.8%",
     "트리거 기반 선택적 로드 — Pareto 최적: 최고 TES + 경쟁력 있는 recall"),
]

SCALABILITY = [
    ("CTSB-small (synthetic)", "50", "12,239", "Yes", "9.8%"),
    ("AgentNode (real)", "215", "409,380", "No (13× 초과)", "3.6%"),
    ("OneViral (real)", "299", "735,309", "No (23× 초과)", "2.7%"),
]


def build_report(out_path: Path):
    doc = Document()

    # ── 페이지 여백 ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══ 표지 ══════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CTX vs Nemotron-Cascade-2")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = C_DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("코드 검색 성능 비교 보고서 (논문급 다각화 평가)")
    r2.font.size = Pt(13); r2.font.color.rgb = C_ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run("날짜: 2026-03-27  |  데이터셋: CTSB-small (166 queries, 50 files)  |  NIPA GPU7 실측")
    r3.font.size = Pt(9); r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60); r3.font.italic = True

    doc.add_paragraph()

    # ══ 0. 요약 ══════════════════════════════════════════════
    add_section_title(doc, "0. 핵심 요약 (Executive Summary)")
    add_body(doc,
        "CTX(adaptive_trigger)와 Nemotron-Cascade-2-30B-A3B(Mamba SSM LLM 랭킹)를 "
        "동일 벤치마크 위에서 비교한 결과, Recall@5는 통계적으로 동등(0.958 vs 0.946, p=0.629)하지만 "
        "CTX의 TES(효율)가 2.78× 우수하고(0.668 vs 0.241, p<1e-27, d=1.32) "
        "Recall@1·NDCG@5도 유의미하게 높습니다. "
        "Nemotron은 소규모 신규 코드베이스의 EXPLICIT_SYMBOL 쿼리에서만 +5pp 우위를 보이며, "
        "실제 프로덕션 규모(40만+ 토큰)에서는 컨텍스트 한계(32K)로 사용 불가합니다."
    )

    # ══ 1. 전체 지표 (설명 포함) ══════════════════════════
    add_section_title(doc, "1. 전체 지표 비교 (설명 포함)")
    add_body(doc,
        "아래 표는 CTX와 Nemotron의 주요 지표를 나열하고, "
        "각 지표의 의미·측정 방법·해석을 설명 열로 제공합니다.",
        italic=True
    )

    # 헤더 행
    cols = ["지표", "CTX", "Nemotron", "Δ (CTX−Nem)", "p-value", "Effect d", "우위", "지표 설명"]
    widths = [Cm(2.4), Cm(1.4), Cm(1.6), Cm(1.6), Cm(2.0), Cm(2.2), Cm(1.4), Cm(6.0)]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = 'Table Grid'

    hdr = tbl.rows[0]
    for i, (col, w) in enumerate(zip(cols, widths)):
        hdr.cells[i].width = w
        set_cell_bg(hdr.cells[i], C_DARK)
        set_cell_font(hdr.cells[i], col, bold=True, size=9,
                      color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, row in enumerate(OVERALL_METRICS):
        tr = tbl.add_row()
        bg = C_GRAY if idx % 2 == 0 else C_WHITE
        winner_col = C_GREEN if row["winner"] == "CTX" else (
                     C_GOLD if row["winner"] == "동등" else C_RED)

        vals = [row["metric"], row["ctx"], row["nem"], row["delta"],
                row["p_value"], row["effect_d"], row["winner"], row["description"]]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]

        for i, (v, a) in enumerate(zip(vals, aligns)):
            set_cell_bg(tr.cells[i], C_LGRAY if i == 6 else bg)
            bold = (i == 0 or i == 6)
            color = winner_col if i == 6 else C_TEXT
            set_cell_font(tr.cells[i], v, bold=bold, size=9, color=color, align=a)

    doc.add_paragraph()

    # ══ 2. 트리거 유형별 분석 ══════════════════════════════
    add_section_title(doc, "2. 트리거 유형별 Recall@5 분석")
    add_body(doc,
        "쿼리를 4가지 트리거 유형으로 분류하여 각 시스템의 강약점을 파악합니다. "
        "설명 열에는 각 트리거 유형의 정의와 결과 해석이 포함됩니다.",
        italic=True
    )

    cols2 = ["트리거 유형", "N", "CTX R@5", "Nem R@5", "Δ", "p-value", "설명"]
    widths2 = [Cm(3.6), Cm(0.8), Cm(1.5), Cm(1.5), Cm(1.3), Cm(1.8), Cm(8.5)]
    tbl2 = doc.add_table(rows=1, cols=len(cols2))
    tbl2.style = 'Table Grid'

    hdr2 = tbl2.rows[0]
    for i, (col, w) in enumerate(zip(cols2, widths2)):
        hdr2.cells[i].width = w
        set_cell_bg(hdr2.cells[i], C_BLUE)
        set_cell_font(hdr2.cells[i], col, bold=True, size=9,
                      color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, (tt, n, ctx_r, nem_r, delta, p_val, desc) in enumerate(TRIGGER_METRICS):
        tr = tbl2.add_row()
        bg = C_GRAY if idx % 2 == 0 else C_WHITE
        try:
            delta_f = float(delta.replace('−', '-'))
        except ValueError:
            delta_f = 0.0
        delta_col = C_GREEN if delta_f > 0 else (C_RED if delta_f < -0.05 else C_GOLD)

        row_vals = [tt, str(n), ctx_r, nem_r, delta, p_val, desc]
        row_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.LEFT]
        for i, (v, a) in enumerate(zip(row_vals, row_aligns)):
            set_cell_bg(tr.cells[i], bg)
            col = delta_col if i == 4 else C_TEXT
            set_cell_font(tr.cells[i], v, bold=(i == 0), size=9, color=col, align=a)

    doc.add_paragraph()

    # ══ 3. 다계층 시스템 비교 ════════════════════════════
    add_section_title(doc, "3. 다계층 시스템 비교 (6개 전략)")
    add_body(doc,
        "Full Context(비랭킹 기준), GraphRAG, RANGER, Nemotron, BM25, CTX 6개 전략을 "
        "동일 벤치마크(n=166)에서 비교합니다. Pareto 최적 시스템 판단 기준 포함.",
        italic=True
    )

    cols3 = ["시스템", "R@1", "R@5", "NDCG@5", "TES", "토큰 사용", "특성 설명"]
    widths3 = [Cm(3.5), Cm(1.2), Cm(1.2), Cm(1.4), Cm(1.2), Cm(1.6), Cm(8.9)]
    tbl3 = doc.add_table(rows=1, cols=len(cols3))
    tbl3.style = 'Table Grid'

    hdr3 = tbl3.rows[0]
    for i, (col, w) in enumerate(zip(cols3, widths3)):
        hdr3.cells[i].width = w
        set_cell_bg(hdr3.cells[i], C_ACCENT)
        set_cell_font(hdr3.cells[i], col, bold=True, size=9,
                      color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    highlight_rows = {"CTX (adaptive_trigger)": C_GREEN, "BM25": RGBColor(0x1A, 0x3A, 0x5C)}
    for idx, row_data in enumerate(MULTI_SYSTEM):
        name = row_data[0]
        tr = tbl3.add_row()
        bg = highlight_rows.get(name, C_GRAY if idx % 2 == 0 else C_WHITE)
        is_highlight = name in highlight_rows
        for i, v in enumerate(row_data):
            set_cell_bg(tr.cells[i], bg)
            txt_col = C_WHITE if is_highlight else C_TEXT
            set_cell_font(tr.cells[i], v, bold=(i == 0 and is_highlight),
                          size=9, color=txt_col,
                          align=WD_ALIGN_PARAGRAPH.CENTER if i in (1,2,3,4,5) else WD_ALIGN_PARAGRAPH.LEFT)

    add_body(doc, "* 녹색 강조: Pareto 최적 (CTX — TES 최고, Recall 경쟁력 유지)  |  청색 강조: Recall@5 최고 (BM25)", italic=True)

    # ══ 4. 확장성 분석 ══════════════════════════════════
    add_section_title(doc, "4. 실코드베이스 확장성 분석")
    add_body(doc,
        "Nemotron(32K 컨텍스트 한계)은 소형 합성 데이터셋에서만 작동합니다. "
        "실제 프로덕션 코드베이스에서는 CTX만 사용 가능합니다.",
        italic=True
    )

    cols4 = ["코드베이스", "파일 수", "총 토큰", "Nemotron 가능?", "CTX 토큰 효율"]
    widths4 = [Cm(4.0), Cm(1.8), Cm(2.8), Cm(3.5), Cm(3.0)]
    tbl4 = doc.add_table(rows=1, cols=len(cols4))
    tbl4.style = 'Table Grid'

    hdr4 = tbl4.rows[0]
    for i, (col, w) in enumerate(zip(cols4, widths4)):
        hdr4.cells[i].width = w
        set_cell_bg(hdr4.cells[i], C_DARK)
        set_cell_font(hdr4.cells[i], col, bold=True, size=9,
                      color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, (name, files, tokens, feasible, eff) in enumerate(SCALABILITY):
        tr = tbl4.add_row()
        bg = C_GRAY if idx % 2 == 0 else C_WHITE
        feas_col = C_GREEN if "Yes" in feasible else C_RED
        for i, (v, col) in enumerate(zip([name, files, tokens, feasible, eff],
                                          [C_TEXT, C_TEXT, C_TEXT, feas_col, C_TEXT])):
            set_cell_bg(tr.cells[i], bg)
            set_cell_font(tr.cells[i], v, bold=(i == 0), size=9, color=col,
                          align=WD_ALIGN_PARAGRAPH.LEFT if i in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ══ 5. 하이브리드 전략 ══════════════════════════════
    add_section_title(doc, "5. 권장 하이브리드 전략")
    add_body(doc,
        "CTX-First + Nemotron-Fallback: CTX를 기본 경로로 사용하고, "
        "소형 신규 코드베이스(<20K 토큰)에서 EXPLICIT_SYMBOL 신뢰도가 낮을 때만 "
        "Nemotron을 fallback으로 호출합니다."
    )

    hybrid_rows = [
        ("활성화 조건", "코드베이스 ≤ 20,000 토큰 AND CTX 신뢰도 < 0.80 AND 트리거=EXPLICIT_SYMBOL"),
        ("예상 Recall@5", "~0.975 (+1.7pp vs CTX, +2.9pp vs Nemotron 단독)"),
        ("예상 TES", "~0.55 (Nemotron은 전체 쿼리의 ~7%에만 활성화)"),
        ("비용-편익 공식", "Nemotron ROI = (0.051 recall gain × query_value) / (10× token cost)"),
        ("권장 적용 케이스", "인덱스 미구축 신규 레포 디버깅, 함수명 모호한 고가치 쿼리"),
    ]

    tbl5 = doc.add_table(rows=len(hybrid_rows), cols=2)
    tbl5.style = 'Table Grid'
    for idx, (label, val) in enumerate(hybrid_rows):
        bg = C_GRAY if idx % 2 == 0 else C_WHITE
        set_cell_bg(tbl5.rows[idx].cells[0], C_LGRAY)
        set_cell_bg(tbl5.rows[idx].cells[1], bg)
        set_cell_font(tbl5.rows[idx].cells[0], label, bold=True, size=9,
                      color=C_ACCENT, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_font(tbl5.rows[idx].cells[1], val, size=9,
                      color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
        tbl5.rows[idx].cells[0].width = Cm(4.5)
        tbl5.rows[idx].cells[1].width = Cm(14.5)

    doc.add_paragraph()

    # ══ 6. 결론 ══════════════════════════════════════
    add_section_title(doc, "6. 결론")
    conclusions = [
        "Recall@5는 통계적으로 동등 (p=0.629) — 커버리지 측면에서 두 시스템은 실질적 차이 없음",
        "CTX는 Recall@1(+16pp***)·NDCG@5(+7.9pp***)·TES(+42.8pp***) 모두 유의미하게 우수",
        "Nemotron은 Pareto frontier 위에 없음 — BM25가 더 낮은 비용에 더 높은 recall 달성",
        "Nemotron의 유일한 가치: 인덱스 불필요, zero-shot, 소형 코드베이스 EXPLICIT_SYMBOL +5pp",
        "실 프로덕션(40만+ 토큰) 스케일에서는 CTX만 물리적으로 작동 가능",
        "권장: CTX를 기본 전략으로 사용, 특수 케이스에만 Nemotron 폴백 적용",
    ]
    for i, c in enumerate(conclusions):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(c)
        run.font.size = Pt(10)
        run.font.color.rgb = C_TEXT

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = foot.add_run(
        "실험: NIPA GPU7 (Nemotron-Cascade-2-30B-A3B, port 8010)  |  "
        "벤치마크: CTSB-small 166 queries  |  "
        "통계: Wilcoxon signed-rank (paired, N=166)"
    )
    r_foot.font.size = Pt(8)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ══ 7. G1/G2 문서 검색 비교 ══════════════════════════════════
    doc.add_page_break()
    add_section_title(doc, "7. Goal 1 / Goal 2 — 문서 검색 비교 실험")

    intro7 = doc.add_paragraph()
    intro7.paragraph_format.space_after = Pt(6)
    r7 = intro7.add_run(
        "CTX의 두 핵심 목표(G1: 연상 기억으로 문서 서페이싱, G2: 지시→유관 파일 검색)를 "
        "동일 Nemotron-Cascade-2 서버로 평가. 벤치마크: CTX docs/ 29개 .md 파일, 87개 쿼리 "
        "(heading_exact / heading_paraphrase / keyword 각 29개). "
        "Nemotron은 29개 파일(각 3000자 트런케이션, ~17K tok) + 쿼리를 단일 프롬프트로 전달해 top-5 순위 결정."
    )
    r7.font.size = Pt(10)
    r7.font.color.rgb = C_TEXT

    # 7-A: Overall 비교표
    add_section_title(doc, "7-A. 전체 문서 검색 성능", level=2)
    hdr7a = ["지표", "설명", "CTX-doc", "Nemotron", "BM25", "Delta (CTX-Nem)", "p-value"]
    rows7a = [
        ("R@3",   "상위 3개 안에 정답 포함",      "0.713", "0.540", "0.667", "+17.3%p", "—"),
        ("R@5",   "상위 5개 안에 정답 포함",      "0.862", "0.586", "0.839", "+27.6%p", "3.0×10⁻⁶ ***"),
        ("NDCG@5","순위 고려 관련성 점수",         "0.717", "0.472", "0.655", "+24.5%p", "—"),
        ("MRR",   "정답이 처음 나오는 순위 역수",  "0.688", "0.433", "0.611", "+25.5%p", "—"),
    ]
    tbl7a = doc.add_table(rows=1 + len(rows7a), cols=7)
    tbl7a.style = 'Table Grid'
    for ci, h in enumerate(hdr7a):
        set_cell_bg(tbl7a.rows[0].cells[ci], C_BLUE)
        set_cell_font(tbl7a.rows[0].cells[ci], h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    col_w7a = [Cm(1.8), Cm(4.5), Cm(1.8), Cm(2.0), Cm(1.8), Cm(2.5), Cm(2.5)]
    for ci, w in enumerate(col_w7a):
        for row in tbl7a.rows:
            row.cells[ci].width = w
    for ri, row_data in enumerate(rows7a):
        row = tbl7a.rows[ri + 1]
        bg = C_LGRAY if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            set_cell_bg(row.cells[ci], bg)
            bold = ci in (2,)  # CTX column bold
            color = C_GREEN if ci == 2 else (C_RED if ci == 3 else C_TEXT)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(row.cells[ci], val, bold=bold, color=color, align=align)

    doc.add_paragraph()

    # 7-B: 쿼리 타입별 (G1/G2 매핑)
    add_section_title(doc, "7-B. 쿼리 타입별 분석 — G1 / G2 매핑", level=2)
    hdr7b = ["쿼리 타입", "CTX 목표", "CTX R@5", "Nemotron R@5", "Delta R@5", "p-value", "해석"]
    rows7b = [
        ("heading_exact",      "G2: 정확 검색", "0.897", "0.655", "+24.1%p", "0.004 **",
         "정확한 헤딩 매칭 — CTX symbol_index 우위"),
        ("heading_paraphrase", "G1: 연상 기억", "1.000", "0.828", "+17.2%p", "0.013 *",
         "G1 핵심: CTX 완벽 달성, Nem LLM 이해로 경쟁"),
        ("keyword",            "G2: 키워드 검색", "0.690", "0.276", "+41.4%p", "0.001 **",
         "keyword: LLM 스니펫 기반 판단 실패, CTX concept_index 압도"),
    ]
    tbl7b = doc.add_table(rows=1 + len(rows7b), cols=7)
    tbl7b.style = 'Table Grid'
    for ci, h in enumerate(hdr7b):
        set_cell_bg(tbl7b.rows[0].cells[ci], C_ACCENT)
        set_cell_font(tbl7b.rows[0].cells[ci], h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    col_w7b = [Cm(2.5), Cm(2.2), Cm(1.8), Cm(2.3), Cm(2.0), Cm(2.0), Cm(4.2)]
    for ci, w in enumerate(col_w7b):
        for row in tbl7b.rows:
            row.cells[ci].width = w
    for ri, row_data in enumerate(rows7b):
        row = tbl7b.rows[ri + 1]
        bg = C_LGRAY if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            set_cell_bg(row.cells[ci], bg)
            bold = ci == 2
            color = C_GREEN if ci == 2 else (C_RED if ci == 3 else C_TEXT)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci in (2, 3, 4, 5) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(row.cells[ci], val, bold=bold, color=color, align=align)

    doc.add_paragraph()

    # 7-C: 코드 + 문서 통합 비교
    add_section_title(doc, "7-C. 코드 + 문서 전영역 통합 비교", level=2)
    hdr7c = ["태스크", "CTX", "Nemotron", "Delta", "p-value", "효과크기"]
    rows7c = [
        ("코드 R@5 (166 queries)",      "0.958", "0.946", "+1.2%p",   "0.629 (ns)", "d=0.042 negligible"),
        ("코드 TES",                    "0.668", "0.241", "+177%",    "3.1×10⁻²⁷ ***", "d=1.322 large"),
        ("문서 R@5 (87 queries)",       "0.862", "0.586", "+27.6%p",  "3.0×10⁻⁶ ***",  "h=0.637 large"),
        ("G1: paraphrase R@5",          "1.000", "0.828", "+17.2%p",  "0.013 *",   "—"),
        ("G2: keyword R@5",             "0.690", "0.276", "+41.4%p",  "0.001 **",  "—"),
        ("G2: 코드 TES (효율성)",       "0.668", "0.241", "+177%",    "p<10⁻²⁶",  "large"),
        ("확장성 (>32K codebase)",      "✓ 지원", "✗ 불가", "독점 강점", "—", "—"),
        ("토큰 사용률",                  "5.2%",  "~100%", "19x 절감", "—", "—"),
    ]
    tbl7c = doc.add_table(rows=1 + len(rows7c), cols=6)
    tbl7c.style = 'Table Grid'
    for ci, h in enumerate(hdr7c):
        set_cell_bg(tbl7c.rows[0].cells[ci], C_DARK)
        set_cell_font(tbl7c.rows[0].cells[ci], h, bold=True, color=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    col_w7c = [Cm(4.0), Cm(2.0), Cm(2.2), Cm(2.2), Cm(2.8), Cm(3.8)]
    for ci, w in enumerate(col_w7c):
        for row in tbl7c.rows:
            row.cells[ci].width = w
    for ri, row_data in enumerate(rows7c):
        row = tbl7c.rows[ri + 1]
        bg = C_LGRAY if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            set_cell_bg(row.cells[ci], bg)
            bold = ci in (0, 1)
            color = C_GREEN if ci == 1 else (C_RED if ci == 2 else C_TEXT)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci != 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(row.cells[ci], val, bold=bold, color=color, align=align)

    doc.add_paragraph()
    foot7 = doc.add_paragraph()
    foot7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot7 = foot7.add_run(
        "G1/G2 실험: NIPA GPU7 (Nemotron-Cascade-2-30B-A3B, port 8010)  |  "
        "CTX docs/ 29 files, 87 queries  |  "
        "통계: Wilcoxon signed-rank (N=87)"
    )
    r_foot7.font.size = Pt(8)
    r_foot7.font.italic = True
    r_foot7.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(out_path))
    print(f"[DONE] {out_path}")


if __name__ == "__main__":
    out = Path(__file__).parent / "CTX_NEMOTRON_COMPARISON_REPORT.docx"
    build_report(out)
