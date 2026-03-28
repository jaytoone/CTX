"""
CTX 종합 실험 보고서 생성기
산출물: docs/research/CTX_FULL_EXPERIMENT_REPORT.docx

내용:
  1. 실험 히스토리 타임라인
  2. 전략별 성능 비교 (CTX vs BM25 vs TF-IDF vs Nemotron)
  3. 쿼리 타입별 세부 성능
  4. 잘하는 것 / 못하는 것
  5. 개선 이력 Before/After
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 팔레트 ────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1A, 0x23, 0x4E)
C_BLUE   = RGBColor(0x1E, 0x40, 0xAF)
C_TEAL   = RGBColor(0x0F, 0x76, 0x6E)
C_GREEN  = RGBColor(0x16, 0x65, 0x34)
C_RED    = RGBColor(0x9B, 0x1C, 0x1C)
C_AMBER  = RGBColor(0x92, 0x40, 0x09)
C_GRAY   = RGBColor(0xF3, 0xF4, 0xF6)
C_LGRAY  = RGBColor(0xE5, 0xE7, 0xEB)
C_HDRBG  = RGBColor(0x1E, 0x40, 0xAF)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT   = RGBColor(0x11, 0x18, 0x27)
C_PALE_G = RGBColor(0xDC, 0xFC, 0xE7)
C_PALE_Y = RGBColor(0xFE, 0xF9, 0xC3)
C_PALE_R = RGBColor(0xFE, 0xE2, 0xE2)
C_PALE_B = RGBColor(0xDB, 0xEA, 0xFE)


def _set_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def _cell(cell, text, bold=False, size=9.5, color=C_TEXT,
          align=WD_ALIGN_PARAGRAPH.CENTER, bg=None):
    if bg:
        _set_bg(cell, bg)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _hdr_row(tbl, cols, widths=None):
    """Add a styled header row."""
    row = tbl.rows[0]
    for i, (cell, col) in enumerate(zip(row.cells, cols)):
        _cell(cell, col, bold=True, size=9.5, color=C_WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER, bg=C_HDRBG)
    if widths:
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w)


def _section(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15 if level == 1 else 12)
    r.font.color.rgb = C_NAVY if level == 1 else C_BLUE


def _body(doc, text, italic=False, size=10):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        r.font.size = Pt(size)
        if italic:
            r.font.italic = True


def _zebra(row_idx):
    return C_GRAY if row_idx % 2 == 0 else C_WHITE


def build_report():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)

    # ── 타이틀 페이지 ─────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run('CTX 종합 실험 보고서')
    title_r.bold = True
    title_r.font.size = Pt(22)
    title_r.font.color.rgb = C_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run(
        'Contextual Trigger eXtraction — BM25 교체 + 쿼리 타입별 라우팅 최적화\n'
        '실험 기간: 2026-03-27  |  최종 업데이트: 2026-03-27'
    )
    sub_r.font.size = Pt(11)
    sub_r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 1. 실험 히스토리
    # ═══════════════════════════════════════════════════════════
    _section(doc, '1. 실험 히스토리')
    _body(doc, '세션 내 수행한 주요 실험 단계 및 각 단계의 결과 요약입니다.')

    history = [
        ('Phase', '실험 내용', '핵심 결과', '상태'),
        ('Phase 1', 'CTX 아키텍처 파악',
         'Rule-based 비LLM 시스템 확인\n4트리거 타입: EXPLICIT_SYMBOL / SEMANTIC_CONCEPT\n/ TEMPORAL_HISTORY / IMPLICIT_CONTEXT',
         '완료'),
        ('Phase 2', 'CTX vs Nemotron-Cascade-2\nG1/G2 비교 실험 (29 docs, 87 queries)',
         'CTX R@5=0.862 vs Nemotron 0.586\n(Wilcoxon p=3.0×10⁻⁶, Cohen h=0.637 대효과)\n코드 TES: CTX 0.668 vs Nemotron 0.241 (+177%)',
         '완료'),
        ('Phase 3', 'CTX 약점 분석 + 대안 조사\n(expert-research)',
         '3대 약점: 외부코드베이스 R@5=0.152,\nkeyword R@3=0.379 < BM25=0.667,\n교차파일 추론 불가',
         '완료'),
        ('Phase 4', 'TF-IDF → BM25 교체 구현\n(AdaptiveTriggerRetriever)',
         'keyword R@3: 0.379 → 0.655 (+72.8%)\n전체 R@3: 0.713 → 0.839 (+17.7%)\n커밋: 5099f32, 7d1a6a8',
         '완료'),
        ('Phase 5\n(시도1)', 'BM25 blend 비율 튜닝 3회\n(norm*0.9 → norm → max+0.05 등)',
         '모두 0.655 plateau\n원인: heading false positive 4건이\nBM25 정답 파일을 밀어냄',
         '수렴(한계)'),
        ('Phase 5\n(최종)', 'query_type-aware routing 구현\nkeyword → TF-only BM25 직접 라우팅\nheading → heading+BM25Okapi',
         'keyword R@3: 0.655 → 0.724 (+10.5%)\n전체 R@3: 0.839 → 0.862 (+2.7%)\nNDCG@5: 0.787 → 0.830\n커밋: f42a22b',
         '✅ 목표달성'),
    ]

    col_w = [0.8, 2.2, 3.2, 0.9]
    tbl = doc.add_table(rows=len(history), cols=4)
    tbl.style = 'Table Grid'
    _hdr_row(tbl, history[0], col_w)
    phase_bg = {
        '완료':    C_PALE_G,
        '수렴(한계)': C_PALE_Y,
        '✅ 목표달성': C_PALE_B,
    }
    for ri, row_data in enumerate(history[1:], 1):
        r = tbl.rows[ri]
        status = row_data[3]
        bg = phase_bg.get(status, C_WHITE)
        _cell(r.cells[0], row_data[0], bold=True, size=9, bg=bg,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(r.cells[1], row_data[1], size=9, bg=bg,
              align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell(r.cells[2], row_data[2], size=9, bg=bg,
              align=WD_ALIGN_PARAGRAPH.LEFT)
        status_color = C_GREEN if '✅' in status else (C_AMBER if '한계' in status else C_TEAL)
        _cell(r.cells[3], status, bold=True, size=9, color=status_color, bg=bg,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 2. 전략별 성능 비교 (최종)
    # ═══════════════════════════════════════════════════════════
    _section(doc, '2. 전략별 성능 비교 (최종, 87 queries · 29 docs)')
    _body(doc, '문서 검색 벤치마크 최종 결과. CTX-doc이 모든 지표에서 우위.')

    compare_data = [
        ('전략', 'R@3', 'R@5', 'NDCG@5', 'MRR', '비고'),
        ('CTX-doc (heading+BM25)',       '0.862', '0.954', '0.830', '0.795', '최고'),
        ('BM25 단독 (TF-only)',          '0.667', '0.839', '0.655', '0.611', 'baseline'),
        ('Dense TF-IDF (cosine)',        '0.690', '0.805', '0.607', '0.563', 'baseline'),
        ('Nemotron-Cascade-2 (G2 문서)', '0.540', '0.586', '—',     '—',     'LLM기반'),
    ]
    cw2 = [2.5, 0.7, 0.7, 0.8, 0.7, 0.8]
    tbl2 = doc.add_table(rows=len(compare_data), cols=6)
    tbl2.style = 'Table Grid'
    _hdr_row(tbl2, compare_data[0], cw2)
    bgs2 = [C_PALE_B, _zebra(1), _zebra(2), _zebra(3)]
    for ri, row_d in enumerate(compare_data[1:], 1):
        r = tbl2.rows[ri]
        bg = bgs2[ri - 1]
        is_ctx = ri == 1
        for ci, val in enumerate(row_d):
            bold_flag = is_ctx or ci == 0
            clr = C_BLUE if (is_ctx and 1 <= ci <= 4) else C_TEXT
            _cell(r.cells[ci], val, bold=bold_flag, size=9.5,
                  color=clr, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 3. 쿼리 타입별 세부 성능
    # ═══════════════════════════════════════════════════════════
    _section(doc, '3. 쿼리 타입별 세부 성능 (R@3)')

    qt_data = [
        ('쿼리 타입', 'N', 'CTX R@3', 'BM25 R@3', 'TF-IDF R@3', 'CTX 우위'),
        ('heading_exact',      '29', '0.862', '0.621', '0.690', '+24.1%p'),
        ('heading_paraphrase', '29', '1.000', '0.655', '0.655', '+34.5%p'),
        ('keyword',            '29', '0.724', '0.724', '0.724', '동등 ✅'),
        ('전체',               '87', '0.862', '0.667', '0.690', '+19.5%p'),
    ]
    cw3 = [1.8, 0.5, 0.9, 0.9, 0.9, 1.1]
    tbl3 = doc.add_table(rows=len(qt_data), cols=6)
    tbl3.style = 'Table Grid'
    _hdr_row(tbl3, qt_data[0], cw3)
    qt_bgs = [C_PALE_G, C_PALE_B, C_PALE_Y, C_GRAY]
    for ri, row_d in enumerate(qt_data[1:], 1):
        r = tbl3.rows[ri]
        bg = qt_bgs[ri - 1]
        for ci, val in enumerate(row_d):
            bold_f = (ci == 0) or (ci == 2)
            clr = C_BLUE if ci == 2 else C_TEXT
            _cell(r.cells[ci], val, bold=bold_f, size=9.5, color=clr, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 4. 코드 검색 성능 (CTX vs Nemotron)
    # ═══════════════════════════════════════════════════════════
    _section(doc, '4. 코드 검색 성능 — CTX vs Nemotron-Cascade-2')
    _body(doc, 'COIR 코드 검색 벤치마크 (real-world 코드베이스 3개: AgentNode, GraphPrompt, OneViral)')

    code_data = [
        ('지표', 'CTX', 'Nemotron', 'Delta', '통계'),
        ('R@5 (전체)',       '0.958', '0.946', '+1.2%p', 'p=0.629 (동등)'),
        ('TES (Token Edit Similarity)', '0.668', '0.241', '+177%', 'p<0.001 ***'),
        ('G2 문서검색 R@3',  '0.839', '0.540', '+55.4%', 'p=3.0×10⁻⁶ ***'),
        ('G2 문서검색 R@5',  '0.931', '0.586', '+58.9%', 'p<0.001 ***'),
    ]
    cw4 = [2.4, 0.9, 0.9, 0.8, 1.3]
    tbl4 = doc.add_table(rows=len(code_data), cols=5)
    tbl4.style = 'Table Grid'
    _hdr_row(tbl4, code_data[0], cw4)
    for ri, row_d in enumerate(code_data[1:], 1):
        r = tbl4.rows[ri]
        bg = _zebra(ri)
        for ci, val in enumerate(row_d):
            is_ctx_val = ci == 1
            clr = C_BLUE if is_ctx_val else C_TEXT
            _cell(r.cells[ci], val, bold=(ci in (0, 1)), size=9.5,
                  color=clr, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 5. 잘하는 것 / 못하는 것
    # ═══════════════════════════════════════════════════════════
    _section(doc, '5. CTX 잘하는 것 / 못하는 것')

    strength_data = [
        ('분류', '항목', '수치', '비고'),
        # 잘하는 것
        ('✅ 강점', 'heading 기반 문서 검색',
         'heading_paraphrase R@3 1.000\nheading_exact R@3 0.862',
         'BM25 대비 +34.5%p\n구조 인식 우위'),
        ('✅ 강점', '코드 심볼 정확도 (TES)',
         'TES 0.668\n(Nemotron 0.241)',
         '+177%\n정확한 파일 위치 반환'),
        ('✅ 강점', '응답 속도',
         '<1ms (결정론적 인덱스)',
         'LLM 호출 없음\n무제한 확장 가능'),
        ('✅ 강점', 'heading+BM25 하이브리드',
         '전체 R@3 0.862\n(BM25 단독 0.667)',
         '두 신호 결합 시너지'),
        # 못하는 것
        ('⚠️ 약점', 'keyword 검색 (BM25 동등 수준)',
         'keyword R@3 0.724\n(BM25 단독과 동일)',
         'heading 정보 없는\n자유 키워드에서 한계'),
        ('⚠️ 약점', '외부 코드베이스 일반화',
         'R@5 = 0.152',
         'heuristic 과적합\nAST 파서 필요'),
        ('⚠️ 약점', '의미 기반 교차 파일 추론',
         '2-hop 한계\n(Import graph만 지원)',
         'multi-hop 추론 불가\nGraphRAG 필요'),
        ('⚠️ 약점', '암묵적 개념 쿼리',
         '"auth 관련 파일"\n등 의미 추론 불가',
         'LLM 없이는 불가\n구조적 한계'),
    ]
    cw5 = [0.9, 2.1, 2.0, 2.0]
    tbl5 = doc.add_table(rows=len(strength_data), cols=4)
    tbl5.style = 'Table Grid'
    _hdr_row(tbl5, strength_data[0], cw5)
    for ri, row_d in enumerate(strength_data[1:], 1):
        r = tbl5.rows[ri]
        is_strength = '✅' in row_d[0]
        bg = C_PALE_G if is_strength else C_PALE_Y
        lbl_color = C_GREEN if is_strength else C_AMBER
        _cell(r.cells[0], row_d[0], bold=True, size=9, color=lbl_color, bg=bg,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        for ci in range(1, 4):
            _cell(r.cells[ci], row_d[ci], size=9, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 6. Before / After 개선 이력
    # ═══════════════════════════════════════════════════════════
    _section(doc, '6. 개선 이력 — Before / After')

    ba_data = [
        ('지표', '초기\n(TF-IDF)', 'BM25 교체\n(Phase 4)', 'Routing 최적화\n(Phase 5 최종)', '전체 개선'),
        ('keyword R@3',   '0.379', '0.655', '0.724', '+90.8%'),
        ('heading_para R@3', '0.966', '1.000', '1.000', '+3.5%'),
        ('heading_exact R@3', '0.793', '0.862', '0.862', '+8.7%'),
        ('전체 R@3',      '0.713', '0.839', '0.862', '+20.9%'),
        ('R@5',           '0.862', '0.931', '0.954', '+10.7%'),
        ('NDCG@5',        '0.717', '0.787', '0.830', '+15.8%'),
        ('MRR',           '0.688', '0.748', '0.795', '+15.6%'),
    ]
    cw6 = [1.5, 1.1, 1.2, 1.5, 1.0]
    tbl6 = doc.add_table(rows=len(ba_data), cols=5)
    tbl6.style = 'Table Grid'
    _hdr_row(tbl6, ba_data[0], cw6)
    for ri, row_d in enumerate(ba_data[1:], 1):
        r = tbl6.rows[ri]
        bg = _zebra(ri)
        for ci, val in enumerate(row_d):
            clr = C_BLUE if ci == 3 else (C_GREEN if ci == 4 else C_TEXT)
            bold_f = ci in (0, 3, 4)
            _cell(r.cells[ci], val, bold=bold_f, size=9.5, color=clr, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # 7. 다음 과제
    # ═══════════════════════════════════════════════════════════
    _section(doc, '7. 잔여 약점 & 다음 과제')

    next_data = [
        ('우선순위', '과제', '현재 수치', '접근법', '예상 효과'),
        ('P1', 'keyword R@3 초과 달성',
         'R@3 = 0.724\n(BM25 동등)',
         'heading+TF-only 조합\nquery expansion',
         'R@3 ≥ 0.800 목표'),
        ('P2', '외부 코드베이스 일반화',
         'R@5 = 0.152',
         'AST 파서 기반 심볼 추출\nheuristic 제거',
         'R@5 ≥ 0.500 목표'),
        ('P3', '교차 파일 추론',
         '2-hop 한계',
         'Import graph BFS 확장\nDependency GraphRAG',
         'multi-hop 추론 가능'),
        ('P4 (장기)', 'LocAgent DHG 통합',
         'CTX: 파일 localization\n평가 미실시',
         'ACL 2025 LocAgent\n(92.7% file-level)',
         '외부 코드베이스\n강점 확보'),
    ]
    cw7 = [0.7, 1.6, 1.3, 1.8, 1.4]
    tbl7 = doc.add_table(rows=len(next_data), cols=5)
    tbl7.style = 'Table Grid'
    _hdr_row(tbl7, next_data[0], cw7)
    pri_colors = {
        'P1': (C_PALE_R, C_RED),
        'P2': (C_PALE_Y, C_AMBER),
        'P3': (C_PALE_G, C_GREEN),
        'P4 (장기)': (C_GRAY, C_TEAL),
    }
    for ri, row_d in enumerate(next_data[1:], 1):
        r = tbl7.rows[ri]
        bg, lbl_c = pri_colors.get(row_d[0], (C_WHITE, C_TEXT))
        _cell(r.cells[0], row_d[0], bold=True, size=9, color=lbl_c, bg=bg,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        for ci in range(1, 5):
            _cell(r.cells[ci], row_d[ci], size=9, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── 푸터 ────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_r = foot_p.add_run(
        'CTX (Contextual Trigger eXtraction)  |  Rule-based 비LLM 검색 시스템  |  2026-03-27'
    )
    foot_r.font.size = Pt(8.5)
    foot_r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    foot_r.font.italic = True

    # ── 저장 ────────────────────────────────────────────────
    out = Path(__file__).parent / 'CTX_FULL_EXPERIMENT_REPORT.docx'
    doc.save(str(out))
    print(f'저장 완료: {out}')


if __name__ == '__main__':
    build_report()
